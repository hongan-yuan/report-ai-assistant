from __future__ import annotations

import csv
from io import StringIO
from pathlib import Path

import chardet
import pdfplumber
from docx import Document
from openpyxl import load_workbook

SUPPORTED_SUFFIXES = {".txt", ".md", ".csv", ".json", ".log", ".pdf", ".docx", ".xlsx", ".xls"}


def _read_text_file(path: Path, max_chars: int) -> str:
    raw = path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"
    text = raw.decode(encoding, errors="replace")
    return text[:max_chars]


def _read_csv(path: Path, max_chars: int) -> str:
    text = _read_text_file(path, max_chars * 2)
    reader = csv.reader(StringIO(text))
    rows = []
    for i, row in enumerate(reader):
        if i > 200:
            rows.append(["... (truncated)"])
            break
        rows.append(row)
    out = "\n".join(",".join(cell for cell in r) for r in rows)
    return out[:max_chars]


def _read_pdf(path: Path, max_chars: int) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:30]:
            parts.append(page.extract_text() or "")
            if sum(len(p) for p in parts) >= max_chars:
                break
    return "\n".join(parts)[:max_chars]


def _read_docx(path: Path, max_chars: int) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)[:max_chars]


def _read_xlsx(path: Path, max_chars: int) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets[:5]:
        parts.append(f"[Sheet: {sheet.title}]")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i > 300:
                parts.append("... (rows truncated)")
                break
            line = "\t".join("" if v is None else str(v) for v in row)
            if line.strip():
                parts.append(line)
    wb.close()
    return "\n".join(parts)[:max_chars]


def extract_text(path: Path, max_chars: int) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".log"}:
        return _read_text_file(path, max_chars)
    if suffix == ".csv":
        return _read_csv(path, max_chars)
    if suffix == ".pdf":
        return _read_pdf(path, max_chars)
    if suffix == ".docx":
        return _read_docx(path, max_chars)
    if suffix in {".xlsx", ".xls"}:
        if suffix == ".xls":
            return "(.xls 格式请先转换为 .xlsx 后分析)"
        return _read_xlsx(path, max_chars)
    return ""


def list_report_files(directory: Path, max_files: int) -> list[Path]:
    if not directory.is_dir():
        return []
    files: list[Path] = []
    for path in sorted(directory.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        files.append(path)
        if len(files) >= max_files:
            break
    return files
